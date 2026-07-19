# -*- coding: utf-8 -*-
"""Bangun Playbook Mitra POS HPY (.docx) dengan screenshot tiap menu."""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION

SHOTS = "docs/screenshots"
OUT = "docs/Playbook-Mitra-POS-HPY.docx"
BLUE = RGBColor(0x1A, 0x73, 0xE8)
GREY = RGBColor(0x5F, 0x63, 0x68)

# (judul, path_gambar, [paragraf deskripsi], [poin bullet])
SECTIONS = [
    ("1. Dashboard", "01-root.png",
     ["Halaman ringkasan setelah login untuk Admin/Manager. Menampilkan indikator "
      "operasional toko secara sekilas."], []),

    ("2. Kasir (POS)", "02-pos.png",
     ["Layar utama penjualan. Tersedia 3 tampilan (Klasik, Quick, Express) yang diatur "
      "admin di Pengaturan Toko."],
     ["Pilih/scan produk → atur qty → (opsional) pilih customer, kupon, nomor meja.",
      "Klik Bayar → pilih metode → isi jumlah dibayar → kembalian dihitung otomatis.",
      "Invoice tersimpan dengan format INV-YYYYMMDD-XXXX; stok berkurang otomatis.",
      "Hold/Recall: tahan pesanan sementara lalu panggil kembali (customer ikut terbawa)."]),

    ("3. Transaksi", "03-transactions.png",
     ["Riwayat seluruh penjualan beserta detail item, pembayaran, dan status sinkron."],
     ["Batalkan transaksi hanya untuk Admin/Manager — stok dikembalikan, data soft-delete."]),

    ("4. Produk", "04-products.png",
     ["Master barang: nama, SKU, harga, kategori, dan opsi lacak stok. Umumnya ditarik "
      "dari ERP HPY (harga Item Price dengan valid_from terbaru)."], []),

    ("5. Customer", "05-customers.png",
     ["Data pelanggan (kode, nama, telepon, poin loyalti). Dapat dipilih saat transaksi "
      "dan didorong (push) ke ERP HPY."], []),

    ("6. Stok Barang", "06-stock.png",
     ["Pantau stok per gudang."],
     ["Sync Bin / Sync Warehouse: tarik saldo stok terbaru dari ERP HPY.",
      "Mendukung banyak gudang; satu gudang ditandai default untuk POS."]),

    ("7. Stock Opname", "07-stockopname.png",
     ["Hitung ulang stok fisik dan cocokkan dengan sistem. Alur: draft → submitted."],
     ["Buat opname, isi qty hasil hitung fisik, lalu Submit. Bisa Cancel saat masih draft."]),

    ("8. Repack (Konversi Item)", "08-slices.png",
     ["Mengubah 1 item menjadi beberapa item lain, mis. 1 Bolu → 8 potong. "
      "Alur: draft → submitted → sync ke ERP HPY (Repack)."],
     ["Submit → stok sumber keluar, hasil masuk.",
      "Catatan: di balik layar masih bernama 'slice', semua label pengguna = Repack."]),

    ("9. Transfer Barang", "09-stocktransfer.png",
     ["Perpindahan stok antar gudang. Outgoing (STO-*) dan Incoming (STI-*)."],
     ["Status: draft → submitted, dengan local_status (draft/sent/received).",
      "Surat Jalan bisa dicetak; tersedia Laporan Transfer."]),

    ("10. Delivery Order", "10-deliveryorders.png",
     ["Pesanan pengiriman ke pelanggan. Nomor DO-YYYYMMDD-XXXX. Alur: draft → confirmed."],
     ["Status dapur: pending → preparing → ready.",
      "Bisa catat pembayaran, atur jadwal produksi, cetak slip Gudang/QC & Invoice.",
      "Sync Sales Order ke ERP HPY."]),

    ("11. Delivery Notes", "11-deliverynotes.png",
     ["Bekerja pada pengiriman (shipment). Tandai Delivered dan Sync Delivery Note ke ERP HPY."], []),

    ("12. Permintaan FG", "12-stockrequests.png",
     ["Permintaan barang jadi ke dapur/pusat. Nomor FG-YYYYMMDD-XXXX. Alur: draft → submitted."],
     ["Status dapur: requested → preparing → done.",
      "Sync ke ERP HPY sebagai Material Request."]),

    ("13. Pulling Order", "13-pullingorder.png",
     ["Layar gabungan untuk mengelola pembayaran dan jadwal produksi lintas Delivery Order "
      "dan Permintaan FG dalam satu tempat."], []),

    ("14. Rekap Order", "14-rekaporder.png",
     ["Laporan agregat kebutuhan produksi per tanggal: gabungan Delivery Order (confirmed, "
      "filter delivery_date) dan Stock Request (submitted, filter needed_date) per kode item."], []),

    ("15. Kitchen Monitor", "15-kitchen.png",
     ["Layar pantau dapur, menampilkan Delivery Order dan Permintaan FG sekaligus dengan "
      "refresh otomatis."],
     ["Ubah status masak hingga ready/done. Tersedia tampilan kalender produksi."]),

    ("16. Sync HPY", "16-sync.png",
     ["Sinkronisasi data dengan server pusat ERP HPY. Berjalan saat itu juga (tanpa antrean)."],
     ["Sync All / per transaksi / Retry Failed untuk mengirim invoice.",
      "Pull: Produk, Metode Pembayaran, User, Harga Delivery, Kupon.",
      "Badge angka = jumlah transaksi belum tersinkron. Status: pending/synced/failed."]),

    ("17. Laporan Online", "17-onlinereport.png",
     ["Data historis penjualan langsung dari ERP HPY."], []),

    ("18. Laporan Pembayaran (MOP)", "18-mopreport.png",
     ["Matriks tanggal × metode pembayaran, ditarik dari ERP HPY."], []),

    ("19. Kupon", "19-coupons.png",
     ["Daftar kupon yang ditarik dari ERP HPY, diterapkan saat checkout di POS."], []),

    ("20. Manajemen User", "20-users.png",
     ["Tambah/ubah user, aktif/nonaktifkan, atur peran dan PIN offline."], []),

    ("21. Manajemen Role", "21-roles.png",
     ["Kelola peran (role) beserta warna penandanya."], []),

    ("22. Hak Akses", "22-permissions.png",
     ["Matriks izin menu per peran. Nyalakan/matikan toggle tiap menu lalu Simpan. "
      "Admin selalu punya akses penuh."],
     ["Modul admin sensitif (Kupon, User, Role, Hak Akses, Warehouse, Pengaturan, Backup, "
      "Update, Factory Reset) default MATI untuk non-admin — nyalakan bila ingin didelegasikan."]),

    ("23. Warehouse", "23-warehouses.png",
     ["Pemetaan gudang ke ERP HPY. Set gudang default (POS) dan transit."], []),

    ("24. Pengaturan Toko", "24-settings.png",
     ["Nama toko, logo, layout POS, pajak/service charge, dan kredensial ERP HPY. "
      "Menu sistem lain: Restore Backup, Update Sistem, Factory Reset (hati-hati, destruktif)."], []),
]


def add_heading(doc, text, size=15, color=BLUE, space_before=14):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return p


def add_para(doc, text, size=10.5, color=None, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    r.italic = italic
    return p


def add_image(doc, path, width_in=6.3):
    if not os.path.exists(path):
        add_para(doc, f"[Screenshot tidak tersedia: {path}]", italic=True, color=GREY)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(path, width=Inches(width_in))


doc = Document()

# Margin lebih rapat agar gambar lega
for s in doc.sections:
    s.top_margin = Inches(0.7)
    s.bottom_margin = Inches(0.7)
    s.left_margin = Inches(0.7)
    s.right_margin = Inches(0.7)

# ── Halaman judul ──
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.paragraph_format.space_before = Pt(120)
r = t.add_run("Playbook Mitra POS HPY")
r.bold = True
r.font.size = Pt(30)
r.font.color.rgb = BLUE

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Panduan Penggunaan Aplikasi Point-of-Sale")
r.font.size = Pt(14)
r.font.color.rgb = GREY

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run("Dilengkapi tangkapan layar tiap menu")
r.font.size = Pt(11)
r.font.color.rgb = GREY
r.italic = True

doc.add_page_break()

# ── Pengantar ──
add_heading(doc, "Tentang Dokumen Ini", size=16)
add_para(doc, "Dokumen ini menjelaskan cara memakai aplikasi Mitra POS HPY untuk "
              "operasional harian: kasir, stok, delivery, produksi dapur, laporan, "
              "hingga pengaturan sistem. Semua integrasi server pusat disebut ERP HPY.")

add_heading(doc, "Login & Peran", size=13, space_before=10)
add_para(doc, "Dua cara masuk: Online (email + password, diverifikasi ke ERP HPY) dan "
              "Offline/PIN (email + PIN 6 digit lokal, dipakai saat server ERP tidak bisa dihubungi).")
add_para(doc, "Landing setelah login: Admin & Manager → Dashboard; Cashier → Kasir (POS); "
              "Dapur → Kitchen Monitor. Menu yang tampil per peran diatur di menu Hak Akses.")

doc.add_page_break()

# ── Isi per menu ──
for title, img, paras, bullets in SECTIONS:
    add_heading(doc, title)
    for ptxt in paras:
        add_para(doc, ptxt)
    for b in bullets:
        bp = doc.add_paragraph(style="List Bullet")
        bp.paragraph_format.space_after = Pt(2)
        br = bp.add_run(b)
        br.font.size = Pt(10)
    add_image(doc, os.path.join(SHOTS, img))
    doc.add_page_break()

# ── Lampiran ──
add_heading(doc, "Lampiran — Kredensial Demo (setelah seeding)", size=13)
tbl = doc.add_table(rows=1, cols=4)
tbl.style = "Light Grid Accent 1"
hdr = tbl.rows[0].cells
for i, h in enumerate(["Peran", "Email", "Password", "PIN"]):
    hdr[i].paragraphs[0].add_run(h).bold = True
for row in [
    ["Admin", "admin@larapos.com", "password", "123456"],
    ["Manager", "manager@larapos.com", "password", "111222"],
    ["Cashier", "kasir@larapos.com", "password", "654321"],
]:
    cells = tbl.add_row().cells
    for i, v in enumerate(row):
        cells[i].text = v

doc.save(OUT)
print("Tersimpan:", OUT)
